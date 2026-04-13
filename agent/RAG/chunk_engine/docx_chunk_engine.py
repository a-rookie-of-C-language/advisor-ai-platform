from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document

from .base_chunk_engine import BaseChunkEngine
from .file_profile import FileProfile


class DocxChunkEngine(BaseChunkEngine):
    name = "docx"

    _DOCX_SUFFIX = {".docx"}
    _DOCX_MIME_KEYWORDS = (
        "wordprocessingml.document",
    )

    def can_handle(self, profile: FileProfile) -> float:
        if profile.extension in self._DOCX_SUFFIX:
            return 1.0
        lowered_mime = (profile.mime or "").lower()
        if any(keyword in lowered_mime for keyword in self._DOCX_MIME_KEYWORDS):
            return 0.95
        return 0.0

    def extract_text(self, file_path: Path) -> str:
        doc = Document(str(file_path))
        lines: List[str] = []

        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                row_text = " | ".join([c for c in cells if c])
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)
